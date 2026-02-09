import io
import PyPDF2
import pdfplumber
from docx import Document
from PIL import Image
import yaml
import json
from typing import Dict, Any, Optional
from pathlib import Path

from app.models import EvidenceType


class DocumentProcessor:
    """Process various document types and extract content"""
    
    @staticmethod
    def process_pdf(file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text and metadata from PDF files"""
        try:
            # Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text = ""
                tables = []
                
                for page in pdf.pages:
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)
                
                metadata = {
                    "num_pages": len(pdf.pages),
                    "has_tables": len(tables) > 0,
                    "table_count": len(tables)
                }
                
                return {
                    "text": text.strip(),
                    "metadata": metadata,
                    "tables": tables,
                    "type": EvidenceType.PDF_DOCUMENT
                }
        except Exception as e:
            # Fallback to PyPDF2
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n\n"
                
                return {
                    "text": text.strip(),
                    "metadata": {"num_pages": len(pdf_reader.pages)},
                    "tables": [],
                    "type": EvidenceType.PDF_DOCUMENT
                }
            except Exception as fallback_error:
                raise Exception(f"PDF processing failed: {str(fallback_error)}")
    
    @staticmethod
    def process_docx(file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from Word documents"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            # Extract text from paragraphs
            text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    table_text.append(row_text)
            
            metadata = {
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
                "num_sections": len(doc.sections)
            }
            
            return {
                "text": text,
                "metadata": metadata,
                "tables": table_text,
                "type": EvidenceType.WORD_DOCUMENT
            }
        except Exception as e:
            raise Exception(f"DOCX processing failed: {str(e)}")
    
    @staticmethod
    def process_image(file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process image files (screenshots, diagrams)"""
        try:
            image = Image.open(io.BytesIO(file_content))
            
            # Detect if it's likely a diagram or screenshot
            width, height = image.size
            is_diagram = any(keyword in filename.lower() 
                           for keyword in ["diagram", "architecture", "network", "topology"])
            
            metadata = {
                "width": width,
                "height": height,
                "format": image.format,
                "mode": image.mode,
                "is_diagram": is_diagram
            }
            
            evidence_type = EvidenceType.NETWORK_DIAGRAM if is_diagram else EvidenceType.SCREENSHOT
            
            # Convert to base64 for Gemini multimodal input (done by caller)
            return {
                "text": None,  # Will be extracted by Gemini vision
                "metadata": metadata,
                "image_data": file_content,
                "type": evidence_type
            }
        except Exception as e:
            raise Exception(f"Image processing failed: {str(e)}")
    
    @staticmethod
    def process_config_file(file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process configuration files (JSON, YAML, etc.)"""
        try:
            content_str = file_content.decode('utf-8')
            
            # Try to parse as JSON or YAML
            parsed_data = None
            file_format = None
            
            if filename.endswith('.json'):
                parsed_data = json.loads(content_str)
                file_format = "json"
            elif filename.endswith(('.yaml', '.yml')):
                parsed_data = yaml.safe_load(content_str)
                file_format = "yaml"
            else:
                # Plain text config
                file_format = "text"
            
            return {
                "text": content_str,
                "metadata": {
                    "format": file_format,
                    "line_count": len(content_str.split('\n'))
                },
                "parsed_data": parsed_data,
                "type": EvidenceType.CONFIG_FILE
            }
        except Exception as e:
            raise Exception(f"Config file processing failed: {str(e)}")
    
    @staticmethod
    def detect_file_type(filename: str, content_type: str) -> EvidenceType:
        """Detect the type of evidence file"""
        filename_lower = filename.lower()
        
        if content_type == "application/pdf" or filename_lower.endswith('.pdf'):
            return EvidenceType.PDF_DOCUMENT
        elif filename_lower.endswith('.docx'):
            return EvidenceType.WORD_DOCUMENT
        elif content_type.startswith('image/'):
            if any(keyword in filename_lower for keyword in ["diagram", "architecture", "network"]):
                return EvidenceType.NETWORK_DIAGRAM
            return EvidenceType.SCREENSHOT
        elif filename_lower.endswith(('.json', '.yaml', '.yml', '.tf', '.hcl')):
            return EvidenceType.CONFIG_FILE
        elif any(keyword in filename_lower for keyword in ["policy", "procedure", "standard"]):
            return EvidenceType.POLICY_DOCUMENT
        else:
            return EvidenceType.UNKNOWN
    
    @staticmethod
    def process_file(file_content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """Main entry point to process any file type"""
        file_type = DocumentProcessor.detect_file_type(filename, content_type)
        
        if file_type == EvidenceType.PDF_DOCUMENT:
            return DocumentProcessor.process_pdf(file_content, filename)
        elif file_type == EvidenceType.WORD_DOCUMENT:
            return DocumentProcessor.process_docx(file_content, filename)
        elif file_type in [EvidenceType.SCREENSHOT, EvidenceType.NETWORK_DIAGRAM]:
            return DocumentProcessor.process_image(file_content, filename)
        elif file_type == EvidenceType.CONFIG_FILE:
            return DocumentProcessor.process_config_file(file_content, filename)
        else:
            # Try to read as text
            try:
                text = file_content.decode('utf-8')
                return {
                    "text": text,
                    "metadata": {},
                    "type": EvidenceType.UNKNOWN
                }
            except:
                raise Exception(f"Unsupported file type: {filename}")
